import streamlit as st
import requests
from PIL import Image
import io
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF  # For PDF generation (simple and lightweight, no extra deps needed)
import json  # For potential API handling

# --- 1. API Config ---
# Hugging Face API for Mistral model (to humanize text)
HF_API_URL = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.3"
HF_TOKEN = st.secrets["HF_TOKEN"]  # Assume you have this in Streamlit secrets
hf_headers = {"Authorization": f"Bearer {HF_TOKEN}"}

# Optional: xAI Grok API (if you want stronger/more advanced humanization)
# Note: You'll need to get an API key from xAI. Replace with your key.
# GROK_API_URL = "https://api.x.ai/v1/chat/completions"  # Hypothetical, check xAI docs for exact endpoint
# GROK_TOKEN = st.secrets.get("GROK_TOKEN", None)  # Optional secret
# If GROK_TOKEN is set, we can use it as fallback or primary for better results.

st.set_page_config(page_title="Professional Resume Builder", layout="wide", initial_sidebar_state="expanded")

# --- 2. Robust Cleaning, Humanization & AI Generation Logic ---
def query_ai(prompt, use_grok=False):
    """
    Query AI (Hugging Face Mistral or Grok) to generate/humanize text.
    Makes output more human-like to bypass detectors.
    """
    try:
        if use_grok and 'GROK_TOKEN' in st.secrets:
            # Use Grok API if available (better for complex prompts)
            grok_headers = {"Authorization": f"Bearer {st.secrets['GROK_TOKEN']}"}
            payload = {
                "model": "grok-beta",  # Adjust based on xAI models
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.8,  # Add randomness for human feel
                "max_tokens": 1024
            }
            response = requests.post(GROK_API_URL, headers=grok_headers, json=payload)
            res = response.json()
            return res['choices'][0]['message']['content'].strip()
        else:
            # Default to Hugging Face
            payload = {"inputs": prompt}
            response = requests.post(HF_API_URL, headers=hf_headers, json=payload)
            res = response.json()
            out = res[0]['generated_text'] if isinstance(res, list) else ""
            # Clean up: Remove prompt echo and tags
            final = out.split(prompt)[-1].strip() if prompt in out else out
            return re.sub(r'<.*?>|\[.*?\]', '', final)  # Remove any AI artifacts
    except Exception as e:
        st.error(f"AI query failed: {str(e)}")
        return ""

def humanize_text(text, label):
    """
    Rewrite text to sound professional and human-like.
    Adds subtle variations: minor typos, varied sentence lengths, personal tone.
    """
    if not text or len(text) < 5: return text
    prompt = f"""
    Rewrite this resume {label} section to sound like a real human wrote it. 
    Make it professional, concise, and natural. 
    Add subtle human touches: varied vocabulary, short/long sentences mix, maybe a small informal phrase if it fits.
    Avoid repetitive patterns. Output plain text only, no markdown or tags.
    Original: {text}
    """
    return query_ai(prompt)

def generate_from_prompt(prompt):
    """
    Generate full resume data from a user prompt using AI.
    AI fills in details, adds some extras if needed.
    """
    gen_prompt = f"""
    Based on this prompt: "{prompt}"
    Generate a complete resume in JSON format with these keys:
    - name: Full name (invent if not specified)
    - job: Job title
    - email: Email
    - phone: Phone number
    - loc: Location
    - sum: About Me summary (200-300 words, human-like)
    - edu: Education details (bullet points or paragraphs)
    - exp: Work Experience (detailed, 3-5 entries, human-written style)
    - skills: Skills list (comma-separated or bullets)
    
    Make it professional, add realistic details if missing, and ensure it sounds human-written (varied language, no AI patterns).
    """
    raw_output = query_ai(gen_prompt, use_grok=True)  # Prefer Grok for generation if available
    try:
        # Parse JSON, clean up
        data = json.loads(raw_output)
        # Humanize each section further
        for key in ['sum', 'edu', 'exp', 'skills']:
            if key in data:
                data[key] = humanize_text(data[key], key)
        return data
    except:
        st.error("Failed to parse AI response. Using defaults.")
        return {}

# --- 3. Document Creation Functions ---
def create_docx(data):
    doc = Document()
    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(data['name'].upper())
    run.font.size = Pt(24)
    run.bold = True
    
    # Job Title
    job_p = doc.add_paragraph()
    job_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    job_p.add_run(data['job']).font.size = Pt(14)
    
    # Contact
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run(f"{data['phone']} | {data['email']} | {data['loc']}").font.size = Pt(10)
    
    # Sections
    for title, content in [("About Me", data['sum']), ("Education", data['edu']), ("Work Experience", data['exp']), ("Skills", data['skills'])]:
        doc.add_heading(title.upper(), level=1)
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 24)
    pdf.cell(200, 10, data['name'].upper(), ln=1, align='C')
    
    pdf.set_font("Arial", "", 14)
    pdf.cell(200, 10, data['job'], ln=1, align='C')
    
    pdf.set_font("Arial", "", 10)
    pdf.cell(200, 10, f"{data['phone']} | {data['email']} | {data['loc']}", ln=1, align='C')
    pdf.ln(10)
    
    for title, content in [("About Me", data['sum']), ("Education", data['edu']), ("Work Experience", data['exp']), ("Skills", data['skills'])]:
        pdf.set_font("Arial", "B", 12)
        pdf.cell(200, 10, title.upper(), ln=1)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 10, content)
        pdf.ln(5)
    
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

# --- 4. UI Layout ---
st.title("📄 Professional AI Resume Builder")
st.markdown("Build a human-like resume that passes AI detectors. Enter details manually or use a prompt for AI generation.")

# Sidebar for options
with st.sidebar:
    st.subheader("⚙️ Settings")
    use_grok = st.checkbox("Use Grok API (if token available)", value=False, help="Grok might give better human-like results. Set GROK_TOKEN in secrets.")
    prompt_mode = st.checkbox("Enable Prompt Mode", value=True, help="Generate from a single prompt.")

# Main columns
col1, col2 = st.columns([1, 1.5])

with col1:
    st.subheader("👤 Input Your Details")
    
    if prompt_mode:
        user_prompt = st.text_area("AI Prompt (e.g., 'Software engineer with 5 years exp in Python')", height=100)
        if st.button("🤖 Generate from Prompt"):
            with st.spinner("Generating resume..."):
                gen_data = generate_from_prompt(user_prompt)
                st.session_state.resume_data = gen_data  # Store in session for editing
    
    # Load data from session or defaults
    resume_data = st.session_state.get("resume_data", {
        "name": "Sebastian Bennett",
        "job": "Professional Accountant",
        "email": "hello@reallygreatsite.com",
        "phone": "+123-456-7890",
        "loc": "Any City, USA",
        "sum": "Experienced accountant specializing in tax...",
        "edu": "Borcelle University | 2026-2030",
        "exp": "Senior Accountant at Salford & Co.",
        "skills": "Auditing, Financial Reporting, Tax Strategy"
    })
    
    # Editable fields (always show for manual input/editing)
    resume_data['name'] = st.text_input("Full Name", resume_data['name'])
    resume_data['job'] = st.text_input("Job Title", resume_data['job'])
    resume_data['email'] = st.text_input("Email", resume_data['email'])
    resume_data['phone'] = st.text_input("Phone", resume_data['phone'])
    resume_data['loc'] = st.text_input("Location", resume_data['loc'])
    resume_data['sum'] = st.text_area("About Me", resume_data['sum'], height=150)
    resume_data['edu'] = st.text_area("Education", resume_data['edu'], height=150)
    resume_data['exp'] = st.text_area("Work Experience", resume_data['exp'], height=200)
    resume_data['skills'] = st.text_area("Skills", resume_data['skills'], height=100)
    
    # Save edits to session
    st.session_state.resume_data = resume_data
    
    if st.button("🚀 Humanize & Generate"):
        with st.spinner("Humanizing content..."):
            # Humanize key sections
            resume_data['sum'] = humanize_text(resume_data['sum'], "summary")
            resume_data['exp'] = humanize_text(resume_data['exp'], "experience")
            resume_data['edu'] = humanize_text(resume_data['edu'], "education")
            resume_data['skills'] = humanize_text(resume_data['skills'], "skills")
            st.session_state.resume_data = resume_data  # Update session

with col2:
    st.subheader("📝 Preview Your Resume")
    if 'resume_data' in st.session_state:
        data = st.session_state.resume_data
        # Nice preview with Markdown (professional styling)
        st.markdown(f"<h1 style='text-align: center; color: #1e3a8a;'>{data['name'].upper()}</h1>", unsafe_allow_html=True)
        st.markdown(f"<h3 style='text-align: center; color: #4b5563;'>{data['job']}</h3>", unsafe_allow_html=True)
        st.markdown(f"<p style='text-align: center; font-size: 14px; color: #6b7280; border-bottom: 1px solid #d1d5db; padding-bottom: 10px;'>{data['phone']} | {data['email']} | {data['loc']}</p>", unsafe_allow_html=True)
        
        st.subheader("ABOUT ME", anchor=False)
        st.markdown(data['sum'])
        
        st.subheader("EDUCATION", anchor=False)
        st.markdown(data['edu'])
        
        st.subheader("WORK EXPERIENCE", anchor=False)
        st.markdown(data['exp'])
        
        st.subheader("SKILLS", anchor=False)
        st.markdown(data['skills'])
        
        # Download options
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            docx_buf = create_docx(data)
            st.download_button("📥 Download as Word (.DOCX)", data=docx_buf, file_name=f"{data['name']}_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col_dl2:
            pdf_buf = create_pdf(data)
            st.download_button("📥 Download as PDF", data=pdf_buf, file_name=f"{data['name']}_Resume.pdf", mime="application/pdf")
    else:
        st.info("Enter details or generate from prompt to preview.")

# Footer
st.markdown("---")
st.caption("Built with Streamlit & AI. Outputs are designed to feel human-written. Edit as needed!")
